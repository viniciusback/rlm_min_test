import streamlit as st
from rlm.rlm_repl import RLM_REPL
from pypdf import PdfReader
from docx import Document
import os
import io
import sys
from contextlib import redirect_stdout
from rlm.logger.streamlit_logger import StreamlitLogHandler
import logging
# =========================

def load_pdf_text(pdf_file) -> str:
    reader = PdfReader(pdf_file)
    text = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text.append(page_text)
    return "\n".join(text)

def streamlit_logger(event):
    if event["type"] == "code_execution":
        log_history.append(
            f"🔁 Iteração {event['iteration']+1}\n"
            f"Executando código:\n```python\n{event['code']}\n```"
        )
        log_box.markdown("\n\n".join(log_history))

# =========================
# Função principal RLM
# =========================
def processar_pdf_com_rlm(context: str, query):
    rlm = RLM_REPL(
        model="gpt-5-mini",
        recursive_model="gpt-5-nano",
        enable_logging=True,
        max_iterations=25
    )

    query = query
    rlm.logger.log = lambda msg: logger.info(msg)
    result = rlm.completion(context=context, query=query)
    return result

# =========================
# Streamlit UI
# =========================
st.set_page_config(page_title="Leitor Jurídico com RLM", layout="wide")

st.title("📑 Agente Leitor - Recursive Language Model")
st.write("Faça upload do PDF")

uploaded_file = st.file_uploader("📤 PDF", type=["pdf"])
query_input = st.text_area("Prompt", height=200)

log_placeholder = st.empty()

logger = logging.getLogger("rlm_streamlit")
logger.setLevel(logging.INFO)

handler = StreamlitLogHandler(log_placeholder)
handler.setFormatter(logging.Formatter("%(message)s"))

if not logger.handlers:   # evita duplicar logs
    logger.addHandler(handler)

# # silenciar lixo de HTTP
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("openai").setLevel(logging.WARNING)

if uploaded_file and query_input:
    log_box = st.empty()

    handler = StreamlitLogHandler(log_box)
    handler.setFormatter(logging.Formatter("%(message)s"))

    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    logger.handlers = []
    logger.addHandler(handler)
    if st.button("⚙️ Processar documento"):
        with st.spinner("Lendo PDF e analisando com RLM... isso pode demorar um pouco 👀"):
            
            # Captura logs (print do RLM)
            log_box = st.empty()
            log_history = []



            # log_buffer = io.StringIO()
            # with redirect_stdout(log_buffer):
            texto_pdf = load_pdf_text(uploaded_file)
            resultado = processar_pdf_com_rlm(texto_pdf,query_input)

            # logs = log_buffer.getvalue()

        st.success("✅ Processo analisado!")

        # =========================
        # Resultado na tela
        # =========================
        st.subheader("📄 Resultado")
        st.text_area("Relatório extraído:", resultado, height=500)

        # =========================
        # Logs escondidos
        # # =========================
        # with st.expander("🪵 Ver logs técnicos (opcional)"):
        #     st.text(logs)

        # =========================
        # Gerar Word
        # =========================
        doc = Document()
        doc.add_heading("Relatório", level=1)

        for linha in resultado.split("\n"):
            doc.add_paragraph(linha)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        st.download_button(
            label="⬇️ Baixar em Word (.docx)",
            data=file_stream,
            file_name="relatorio.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
